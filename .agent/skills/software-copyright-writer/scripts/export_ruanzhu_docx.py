#!/usr/bin/env python3
"""把 Markdown 草稿导出为 .docx。"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path


def ensure_docx_module() -> None:
    try:
        import docx  # noqa: F401

        return
    except Exception:
        pass

    fallback_python = os.environ.get("SOFTWARE_COPYRIGHT_WRITER_PYTHON")
    if (
        fallback_python
        and Path(fallback_python).exists()
        and Path(sys.executable).resolve() != Path(fallback_python).resolve()
    ):
        env = os.environ.copy()
        env["SOFTWARE_COPYRIGHT_WRITER_DOCX_BOOTSTRAPPED"] = "1"
        completed = subprocess.run([fallback_python, __file__, *sys.argv[1:]], env=env, check=False)
        raise SystemExit(completed.returncode)

    raise SystemExit(
        "未找到 python-docx，请先安装 python-docx，或通过 SOFTWARE_COPYRIGHT_WRITER_PYTHON 指定可用 Python。"
    )


ensure_docx_module()

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

HEADING_RE = re.compile(r"^\s*(#{1,6})\s+(.*)$")
HEADING_NUMBER_RE = re.compile(r"^(\d+(?:\.\d+)*)(?:[.、])?\s+.+$")
UNORDERED_RE = re.compile(r"^[-*]\s+(.*)$")
ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
IMAGE_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
PLACEHOLDER_TEXT = "[此处添加对应图片]"
SOURCE_FILE_MARKER_RE = re.compile(r"^===== 文件名:\s*(.*?)\s*=====$")
CODE_FENCE_RE = re.compile(r"^```")
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}


@dataclass
class Block:
    kind: str
    lines: list[str]


@dataclass
class EvidenceImage:
    path: Path
    caption: str


def normalize_markdown_line(raw: str) -> str:
    return raw.replace("\t", "    ").rstrip()


def normalize_heading_line(raw: str) -> str:
    line = normalize_markdown_line(raw)
    heading = HEADING_RE.match(line)
    if heading:
        return f"{heading.group(1)} {heading.group(2).strip()}"
    return line


def set_run_fonts(
    run, chinese_font: str, western_font: str, size_pt: int, bold: bool = False
) -> None:
    run.font.name = western_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)
    run.font.size = Pt(size_pt)
    run.bold = bold


def style_paragraph(
    paragraph,
    chinese_font: str = "宋体",
    western_font: str = "Times New Roman",
    size_pt: int = 12,
    bold: bool = False,
) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, chinese_font, western_font, size_pt, bold=bold)


def apply_heading_style(paragraph, level: int) -> None:
    if level <= 1:
        size = 16
    elif level == 2:
        size = 14
    else:
        size = 12
    style_paragraph(
        paragraph, chinese_font="黑体", western_font="Times New Roman", size_pt=size, bold=True
    )


def style_source_paragraph(paragraph, bold: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, chinese_font="宋体", western_font="Courier New", size_pt=10, bold=bold)


def style_code_paragraph(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    for run in paragraph.runs:
        set_run_fonts(run, chinese_font="宋体", western_font="Courier New", size_pt=10, bold=False)


def style_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph)


def resolve_image_path(raw_path: str, base_dir: Path) -> Path:
    image_path = Path(raw_path.strip()).expanduser()
    if image_path.is_absolute():
        return image_path
    candidate = (base_dir / image_path).resolve()
    if candidate.exists():
        return candidate
    return image_path.resolve()


def add_picture(
    document: Document, image_path: Path, caption: str = "", width_cm: float = 15.5
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph.add_run(caption)
        style_paragraph(caption_paragraph)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="把软著 Markdown 导出为 .docx")
    parser.add_argument("--content", required=True, help="主题内容 Markdown 文件")
    parser.add_argument("--source", required=True, help="源码信息 Markdown 文件")
    parser.add_argument("--output", required=True, help="输出目录")
    parser.add_argument(
        "--evidence-manifest",
        action="append",
        default=[],
        help="fetch_web_evidence.py 生成的 manifest.json，可重复传入；图片会按顺序替换正文占位符",
    )
    parser.add_argument(
        "--image-dir",
        action="append",
        default=[],
        help="额外图片目录，按文件名排序后用于替换正文占位符，可重复传入",
    )
    parser.add_argument(
        "--image-width-cm", type=float, default=15.5, help="正文图片宽度，默认 15.5cm"
    )
    return parser.parse_args()


def load_evidence_images(manifest_paths: list[str], image_dirs: list[str]) -> list[EvidenceImage]:
    images: list[EvidenceImage] = []
    seen: set[Path] = set()

    for raw_manifest_path in manifest_paths:
        manifest_path = Path(raw_manifest_path).expanduser().resolve()
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        for item in manifest:
            screenshot = item.get("screenshot")
            if not screenshot:
                continue
            image_path = (manifest_path.parent / screenshot).resolve()
            if image_path.exists() and image_path not in seen:
                caption = item.get("name") or item.get("title") or image_path.stem
                images.append(EvidenceImage(path=image_path, caption=caption))
                seen.add(image_path)

    for raw_image_dir in image_dirs:
        image_dir = Path(raw_image_dir).expanduser().resolve()
        if not image_dir.exists():
            continue
        for image_path in sorted(image_dir.iterdir()):
            if image_path.suffix.lower() not in IMAGE_EXTENSIONS:
                continue
            resolved = image_path.resolve()
            if resolved not in seen:
                images.append(EvidenceImage(path=resolved, caption=resolved.stem))
                seen.add(resolved)

    return images


def chunk_blocks(lines: list[str]) -> list[Block]:
    blocks: list[Block] = []
    buffer: list[str] = []
    in_table = False
    in_code = False
    code_buffer: list[str] = []

    def flush(kind: str = "paragraph") -> None:
        nonlocal buffer
        if buffer:
            blocks.append(Block(kind=kind, lines=buffer[:]))
            buffer = []

    for raw in lines:
        line = normalize_markdown_line(raw)
        if CODE_FENCE_RE.match(line):
            if in_code:
                blocks.append(Block(kind="code", lines=code_buffer[:]))
                code_buffer = []
                in_code = False
            else:
                flush()
                if in_table:
                    flush(kind="table")
                    in_table = False
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if line.startswith("|") and line.endswith("|"):
            if buffer and not in_table:
                flush()
            in_table = True
            buffer.append(line)
            continue
        if in_table:
            flush(kind="table")
            in_table = False
        if not line.strip():
            flush()
            continue
        if HEADING_RE.match(line):
            flush()
            blocks.append(Block(kind="heading", lines=[normalize_heading_line(line)]))
            continue
        if UNORDERED_RE.match(line):
            flush()
            blocks.append(Block(kind="bullet", lines=[line]))
            continue
        if ORDERED_RE.match(line):
            flush()
            blocks.append(Block(kind="number", lines=[line]))
            continue
        buffer.append(line)
    if in_table:
        flush(kind="table")
    else:
        flush()
    if in_code and code_buffer:
        blocks.append(Block(kind="code", lines=code_buffer[:]))
    return blocks


def validate_heading_hierarchy(lines: list[str]) -> None:
    seen_numbers: set[str] = set()
    next_index_by_parent: dict[str, int] = {"": 1}
    in_code = False

    for line_number, raw in enumerate(lines, start=1):
        line = normalize_markdown_line(raw)
        if CODE_FENCE_RE.match(line):
            in_code = not in_code
            continue
        if in_code:
            continue

        heading = HEADING_RE.match(line)
        if not heading:
            continue

        level = len(heading.group(1))
        text = heading.group(2).strip()
        if level == 1 and not HEADING_NUMBER_RE.match(text):
            continue

        number_match = HEADING_NUMBER_RE.match(text)
        if not number_match:
            raise SystemExit(f"标题层级校验失败：第 {line_number} 行标题缺少编号：{text}")

        number = number_match.group(1)
        parts = number.split(".")
        expected_depth = level - 1
        if len(parts) != expected_depth:
            raise SystemExit(
                f"标题层级校验失败：第 {line_number} 行 `{text}` 应使用 {expected_depth} 级编号，当前是 {len(parts)} 级编号。"
            )

        parent = ".".join(parts[:-1])
        if parent and parent not in seen_numbers:
            raise SystemExit(
                f"标题层级校验失败：第 {line_number} 行 `{text}` 缺少父标题 `{parent}`。"
            )

        current_index = int(parts[-1])
        expected_index = next_index_by_parent.get(parent, 1)
        if current_index != expected_index:
            expected_number = f"{parent}.{expected_index}" if parent else str(expected_index)
            raise SystemExit(
                f"标题层级校验失败：第 {line_number} 行 `{text}` 编号不连续，应为 `{expected_number}`。"
            )

        seen_numbers.add(number)
        next_index_by_parent[parent] = current_index + 1
        next_index_by_parent.setdefault(number, 1)


def add_markdown_paragraph(
    document: Document,
    text: str,
    style: str | None = None,
    base_dir: Path | None = None,
    evidence_images: list[EvidenceImage] | None = None,
    image_width_cm: float = 15.5,
) -> None:
    base_dir = base_dir or Path.cwd()
    image_match = IMAGE_RE.fullmatch(text.strip())
    if image_match:
        image_path = resolve_image_path(image_match.group(2), base_dir)
        if image_path.exists():
            add_picture(document, image_path, image_match.group(1), width_cm=image_width_cm)
        else:
            placeholder = document.add_paragraph()
            placeholder.add_run(PLACEHOLDER_TEXT)
            style_paragraph(placeholder)
        return
    if PLACEHOLDER_TEXT in text:
        prefix = text.replace(PLACEHOLDER_TEXT, "").strip()
        if prefix:
            paragraph = document.add_paragraph(style=style)
            paragraph.add_run(prefix)
            style_paragraph(paragraph)
        if evidence_images:
            evidence = evidence_images.pop(0)
            add_picture(document, evidence.path, evidence.caption, width_cm=image_width_cm)
        else:
            placeholder = document.add_paragraph(style=style)
            placeholder.add_run(PLACEHOLDER_TEXT)
            style_paragraph(placeholder)
        return
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    style_paragraph(paragraph)


def render_table(document: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = value
    style_table(table)


def render_markdown_to_docx(
    markdown_path: Path,
    output_path: Path,
    evidence_images: list[EvidenceImage] | None = None,
    image_width_cm: float = 15.5,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.start_type = WD_SECTION_START.NEW_PAGE

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    validate_heading_hierarchy(lines)
    base_dir = markdown_path.parent
    for block in chunk_blocks(lines):
        if block.kind == "heading":
            match = HEADING_RE.match(block.lines[0])
            level = min(len(match.group(1)), 4)
            heading = document.add_heading(match.group(2).strip(), level=level)
            apply_heading_style(heading, level)
            continue
        if block.kind == "bullet":
            text = UNORDERED_RE.match(block.lines[0]).group(1).strip()
            add_markdown_paragraph(
                document,
                text,
                style="List Bullet",
                base_dir=base_dir,
                evidence_images=evidence_images,
                image_width_cm=image_width_cm,
            )
            continue
        if block.kind == "number":
            text = ORDERED_RE.match(block.lines[0]).group(1).strip()
            add_markdown_paragraph(
                document,
                text,
                style="List Number",
                base_dir=base_dir,
                evidence_images=evidence_images,
                image_width_cm=image_width_cm,
            )
            continue
        if block.kind == "table":
            render_table(document, block.lines)
            continue
        if block.kind == "code":
            for code_line in block.lines:
                paragraph = document.add_paragraph()
                paragraph.add_run(code_line)
                style_code_paragraph(paragraph)
            continue
        paragraph_text = " ".join(line.strip() for line in block.lines).strip()
        if paragraph_text:
            add_markdown_paragraph(
                document,
                paragraph_text,
                base_dir=base_dir,
                evidence_images=evidence_images,
                image_width_cm=image_width_cm,
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def render_source_to_docx(source_path: Path, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.start_type = WD_SECTION_START.NEW_PAGE

    for raw_line in source_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        marker = SOURCE_FILE_MARKER_RE.match(raw_line)
        if marker:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"文件名：{marker.group(1)}")
            style_source_paragraph(paragraph, bold=True)
            continue
        paragraph = document.add_paragraph()
        paragraph.add_run(raw_line)
        style_source_paragraph(paragraph)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> None:
    args = parse_args()
    output_dir = Path(args.output).resolve()
    content_path = Path(args.content).resolve()
    source_path = Path(args.source).resolve()

    if not content_path.exists():
        raise SystemExit(f"未找到主题内容文件：{content_path}")
    if not source_path.exists():
        raise SystemExit(f"未找到源码信息文件：{source_path}")

    content_docx = output_dir / f"{content_path.stem}.docx"
    source_docx = output_dir / f"{source_path.stem}.docx"
    evidence_images = load_evidence_images(args.evidence_manifest, args.image_dir)

    render_markdown_to_docx(
        content_path,
        content_docx,
        evidence_images=evidence_images,
        image_width_cm=args.image_width_cm,
    )
    render_source_to_docx(source_path, source_docx)

    print(f"已导出：{content_docx}")
    print(f"已导出：{source_docx}")


if __name__ == "__main__":
    main()
